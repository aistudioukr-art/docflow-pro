from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
import os, zipfile

app = FastAPI()
app.mount("/", StaticFiles(directory="static", html=True), name="static")

OUTPUT_DIR = "output"
TEMPLATE = "template.docx"

@app.post("/generate")
async def generate(request: Request):
    data = await request.json()

    sender = data["sender"]
    consignee = data["consignee"]
    rows = data["rows"]

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    zip_path = f"{OUTPUT_DIR}/result.zip"
    zipf = zipfile.ZipFile(zip_path, 'w')

    for row in rows:
        doc = Document(TEMPLATE)

        replace = {
            "CONTAINER": row["container"],
            "WEIGHT": row["weight"],
            "EXPIRY": row["expiry"],
            "SENDER": sender,
            "CONSIGNEE": consignee
        }

        for p in doc.paragraphs:
            for k, v in replace.items():
                if k in p.text:
                    p.text = p.text.replace(k, v)

        for table in doc.tables:
            for r in table.rows:
                for cell in r.cells:
                    for k, v in replace.items():
                        if k in cell.text:
                            cell.text = cell.text.replace(k, v)

        filename = f"{OUTPUT_DIR}/cert_{row['container']}.docx"
        doc.save(filename)

        zipf.write(filename)

    zipf.close()

    return FileResponse(zip_path, filename="certificates.zip")